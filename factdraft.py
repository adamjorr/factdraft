#!/usr/bin/env python

from pyzotero import zotero as zotero
import cPickle as pickle
import ConfigParser
import json
import codecs
import re
import os
import sys
import tempfile
import pprint
from subprocess import call, Popen
from docx import Document

NOTEPATH = './draftw_notes'
DECKPATH = './draftw_decks'

def parse_config(config_file):
    config = ConfigParser.ConfigParser()
    config.read(config_file)
    zotero_id = config.get('default','ZOTERO_ID')
    zotero_key = config.get('default','ZOTERO_KEY')
    zotero_type = config.get('default','ZOTERO_TYPE')
    return zotero_id, zotero_type, zotero_key

def request_login_info():
    uid = raw_input("\nWhat is your Zotero API ID? ")
    ktype = raw_input("\nAre you a user or a group? ")
    apikey = raw_input("\nWhat is your Zotero API key? ")
    info = [uid, ktype, apikey]
    return info

def get_coll_key_from_name(zot):
    #get the key of a collection from its name
    #this looks like it doesnt work
    collections = zot.collections()
    string_cols = []
    for collection in collections:
        string_cols.append(collection['data']['name'])
    print "Your collections are:\n" + ' '.join(string_cols)
    collname = raw_input("\n\nWhat collection would you like to use? ")
    for collection in collections:
        if collection['data']['name'] == collname:
            return collection['data']['key']

def request_style_info():
    style = raw_input("\nWhat style would you like to use? ")
    return style

def login(login_info):
    return zotero.Zotero(login_info[0], login_info[1], login_info[2])

def save_obj(obj, filename):
    with open(filename, 'wb') as out:
        pickle.dump(obj,out,-1)

def load_obj(filename):
    with open(filename, 'rb') as input:
        return pickle.load(input)

# def get_collections(zot):
#     for collection in zot.collections():



# class Notecard:
#     def __init__(self, note=''):
#         self.note = note
#         self.source =

def cleanup_html(html_string, i):
    #trends-in-biotechnology is broken so i need to clean up the xml.
    #while we're at it, let's add numbering to the sources.
    cleaned = re.sub(r'<div class="csl-entry" style="clear: left;">\s+.+?</div>',r'',html_string)
    number_search = r'<div class="csl-right-inline" style="margin: 0 .4em 0 1.5em;">'
    numbered = re.sub(number_search, "\n" + str(i) + '. ', cleaned)
    numbered = re.sub(r'</div>\s*?</div>',r'',numbered)
    numbered = re.sub(r'\d\d' + unichr(8211) + r'.+\(','(',numbered) 
    return numbered

def find_itallics(str_with_itallics):
    #find itallics and break up the string
    first_split = str_with_itallics.split('<i>')
    final_split = []
    for my_string in first_split:
        second_split = my_string.split('</i>')
        final_split.extend(second_split)
    return final_split

def make_bibliography_docx(name, bibliography_text):
    #opens and writes to a docx for the bibliography.
    document = Document()
    bibliography = document.add_paragraph()
    all_text = find_itallics(bibliography_text)
    for i in range(0,len(all_text)):
        if i%2==0:
            #is even
            bibliography.add_run(all_text[i]).italic = False
        else:
            bibliography.add_run(all_text[i]).italic = True    
    document.save(name)

def get_bibliography_items(zot, itemIDs, style = None):
    #gets the cleaned html text suitable for passing to make_bibliography_docx.
    #each item will be cited in the order it appears in items.
    style = style or request_style_info()
    cleaned_html = ''
    zot.add_parameters(content = 'bib', style = style)
    items = zot.get_subset(itemIDs)
    for item in items:
        cleaned_html += (cleanup_html(item, items.index(item) + 1))
    return cleaned_html

def open_in_editor(filename):
    EDITOR = os.environ.get('EDITOR','vim')
    call([EDITOR, filename])

# def open_temp_file(zot, itemID):
#     with tempfile.NamedTemporaryFile(suffix=".pdf") as tfile:
#         tfile.write(zot.file(itemID))
#         tfile.flush()
#         Popen(['google-chrome', tfile.name],shell=False,stdin=None,stdout=None,stderr=None,close_fds=True)

def get_items_from_collection(zot, collectionID):
    items = zot.collection_items(collectionID)
    return items

def open_notes_and_files(zot, all_items):
    if not os.path.exists(NOTEPATH): os.mkdir(NOTEPATH)
    pdf_items = []
    for item in all_items:
        if 'contentType' in item['data'] and item['data']['contentType'] == 'application/pdf':
            pdf_items.append(item)
    for item in pdf_items:
        key = item['data']['key']
        parent = item['data']['parentItem']
        fname = NOTEPATH + '/' + parent + '.txt'
        if os.path.exists(fname): continue                      #If the user reopens the program, we don't want them to have to redo everything.
        with tempfile.NamedTemporaryFile(suffix=".pdf") as tfile:
            tfile.write(zot.file(key))
            tfile.flush()
            output = '>/dev/null'
            call(['google-chrome', tfile.name])
            open_in_editor(fname)

def parse_notes():
    #parses every note and returns a dictionary with each heading,
    #then a list of contents
    #and matching list of item ids.
    note_content = {}
    for filename in os.listdir(NOTEPATH):
        item_id = filename.split('.')[0]
        with open(NOTEPATH + '/' + filename,'rb') as f:
            for line in f:
                if len(line.strip().split('\t')) == 2:
                    [content, category] = line.strip().split('\t')
                else:
                    content = line.strip()
                    category = 'uncategorized'
                if not note_content.has_key(category):
                    note_content[category] = {'content' : [content], 'item_id' : [item_id]}
                else:
                    note_content[category]['content'].append(content)
                    note_content[category]['item_id'].append(item_id)
    return note_content

def create_decks(note_contents):
    if not os.path.exists(DECKPATH): os.mkdir(DECKPATH)
    for category in note_contents:
        with open( DECKPATH + "/" + category + ".txt", 'wb') as fo:
            for i in range(0,len(note_contents[category]['content'])):
                new_line = note_contents[category]['content'][i] + '\t' + note_contents[category]['item_id'][i] + '\n'
                fo.write(new_line)

def rearrange_within_decks():
    for filename in os.listdir(DECKPATH):
        if filename.split('.')[len(filename.split('.'))-1] == 'draftw': continue                 #Skip if we've already done this when running the program a different time
        fullname = DECKPATH + "/" + filename
        open_in_editor(fullname)
        os.rename(fullname, fullname + '.draftw')         #We'll use this extension to check if they're done or not.

def reorder_decks():
    deck_names = os.listdir(DECKPATH)
    for i in deck_names: i = i.split('.')[0]
    joined_names = ' '.join(deck_names)
    print "\nThe decks you have are:\n" + joined_names + "\n"
    desired_order_str = raw_input("What order would you like them to be in?\n")
    desired_order_list = desired_order_str.split(' ')
    for deck in desired_order_str.split(' '):
        print deck
        if not deck in deck_names:
            print "There was a problem."
            return reorder_decks()
    for i in desired_order_list: i = i + '.draftw'
    return desired_order_list

def parse_decks(ordered_decks):
    resultdict = {}
    for filename in ordered_decks:
        head_name = filename.split('.')[0]
        section_content = []
        section_ids = []
        with open(DECKPATH + '/' + filename,'rb') as f:
            for line in f:
                if len(line.strip().split('\t')) == 2:
                    [content, item_id] = line.strip().split('\t')
                    section_content.append(content)
                    section_ids.append(item_id)
        resultdict[head_name] = {'content': section_content, 'item_ids': section_ids}
    return resultdict

def get_unique_refs_from_dict(resultdict, desired_order_list):
    item_ids = []
    for headingb in desired_order_list:
        heading = headingb.split('.')[0]
        item_ids.extend(resultdict[heading]['item_ids'])
    unique_item_ids = get_unique_preserve_order(item_ids)
    return unique_item_ids

def get_unique_preserve_order(anylist):
    unique_list = []
    for i in range(0,len(anylist)):
        string = anylist[i]
        if not string in unique_list:
            unique_list.append(string)
    return unique_list

def add_numbered_endings(content_dict,unique_item_ids):
    for heading in content_dict:
        for i in range(0,len(content_dict[heading]['content'])):
            content_dict[heading]['content'][i] = (
                content_dict[heading]['content'][i][:-1] +
                ' [' + str(unique_item_ids.index(content_dict[heading]['item_ids'][i]) + 1) + ']. ')

def clobber_content_dict(content_dict, desired_order_list):
    clobbered = ''
    for heading in desired_order_list:
        heading = heading.split('.')[0]
        clobbered += ('\n\n' + heading + '\n' + '='*len(heading) + '\n\n')
        for content in content_dict[heading]['content']:
            clobbered += content
    clobbered += '\n'
    return clobbered

def main():
    #Login
    zot = login(login_info = [myid, mytype, mykey])


    #Create notes
    items = get_items_from_collection(zot, get_coll_key_from_name(zot))
    open_notes_and_files(zot,items)

    #Create decks
    note_contents = parse_notes()
    create_decks(note_contents)
    rearrange_within_decks()

    #Order decks
    desired_order = reorder_decks()
    content_dict = parse_decks(desired_order)
    unique_item_ids = get_unique_refs_from_dict(content_dict, desired_order)
    add_numbered_endings(content_dict, unique_item_ids)
    draft_content = clobber_content_dict(content_dict, desired_order)
    with open('draftw_out.txt','wb') as f:
        f.write(draft_content)
    bibliography_text = get_bibliography_items(zot, unique_item_ids, style=mystyle_name)
    make_bibliography_docx("references.docx",bibliography_text)

if __name__ == "__main__":
    main()
